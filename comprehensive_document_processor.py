#!/usr/bin/env python3

import os
import json
import re
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
import hashlib

# File processing imports
# Use specific libraries instead of textract which has dependency issues
# import textract  # For various file formats
import PyPDF2
import pdfplumber
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

# Embedding and vector imports
import pandas as pd
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ComprehensiveDocumentProcessor:
    """
    Advanced document processor for legal case files
    Handles multiple file formats and extracts legal-specific metadata
    """
    
    def __init__(self, base_path: str = "/Users/homebase/Desktop/1R-MASTER/DIV-5"):
        self.base_path = Path(base_path)
        self.supported_formats = {
            '.txt': self._process_text,
            '.pdf': self._process_pdf,
            '.json': self._process_json,
            '.docx': self._process_docx,
            '.doc': self._process_doc_fallback,
            '.pages': self._process_pages,
            '.rtfd': self._process_rtfd,
            '.rtf': self._process_rtf
        }
        
        # Legal pattern recognition
        self.legal_patterns = {
            'case_citations': r'\b\w+\s+v\.\s+\w+.*?\d+\s+[A-Za-z\.]+\d*\s+\d+',
            'statute_citations': r'\b\d+\s+Pa\.C\.S\.\s*§\s*\d+[\.\d]*',
            'court_rules': r'\bPa\.R\.C\.P\.\s*\d+[\.\d]*',
            'legal_doctrines': r'(abandonment|doctrine|precedent|holding|dicta)',
            'dates': r'\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b|\b\d{4}-\d{2}-\d{2}\b',
            'dollar_amounts': r'\$[\d,]+\.?\d*',
            'court_names': r'(Court of Common Pleas|Superior Court|Supreme Court|Delaware County)',
        }
        
        # Legal claim patterns for fact-checking
        self.legal_claim_patterns = [
            r'abandonment doctrine',
            r'financial abandonment',
            r'medical abandonment',
            r'precedent.*v\.\s*\w+',
            r'\d+\s+Pa\.C\.S\.\s*§\s*\d+',
            r'Pa\.R\.C\.P\.\s*\d+'
        ]
        
        # Initialize sentence transformer for embeddings
        try:
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("✅ Sentence transformer model loaded")
        except Exception as e:
            logger.warning(f"⚠️ Could not load sentence transformer: {e}")
            self.sentence_model = None
        
        # Document database
        self.processed_documents = []
        self.document_embeddings = {}
        
    def scan_all_directories(self) -> List[Dict]:
        """
        Scan all directories and subdirectories for supported file types
        Returns list of file paths with metadata
        """
        found_files = []
        
        logger.info(f"🔍 Scanning directory: {self.base_path}")
        
        for root, dirs, files in os.walk(self.base_path):
            # Skip hidden directories and specific unwanted directories
            dirs[:] = [d for d in dirs if not d.startswith('.') and d not in ['__pycache__', 'node_modules']]
            
            for file in files:
                file_path = Path(root) / file
                file_ext = file_path.suffix.lower()
                
                if file_ext in self.supported_formats:
                    # Extract directory category from path
                    relative_path = file_path.relative_to(self.base_path)
                    category = self._extract_category_from_path(relative_path)
                    
                    file_info = {
                        'file_path': str(file_path),
                        'relative_path': str(relative_path),
                        'filename': file,
                        'extension': file_ext,
                        'category': category,
                        'size_bytes': file_path.stat().st_size if file_path.exists() else 0,
                        'modified_time': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat() if file_path.exists() else None,
                        'file_hash': self._calculate_file_hash(file_path) if file_path.exists() else None
                    }
                    
                    found_files.append(file_info)
        
        logger.info(f"📁 Found {len(found_files)} supported files across all directories")
        return found_files
    
    def process_all_documents(self, max_files: Optional[int] = None) -> List[Dict]:
        """
        Process all documents in the directory structure
        """
        logger.info("🚀 Starting comprehensive document processing...")
        
        # Get all files
        all_files = self.scan_all_directories()
        
        if max_files:
            all_files = all_files[:max_files]
            logger.info(f"📊 Processing first {max_files} files for testing")
        
        processed_docs = []
        failed_files = []
        
        for i, file_info in enumerate(all_files, 1):
            try:
                logger.info(f"📄 Processing {i}/{len(all_files)}: {file_info['filename']}")
                
                # Process the file
                doc_data = self.process_single_file(file_info['file_path'])
                
                if doc_data:
                    # Add file metadata
                    doc_data.update(file_info)
                    processed_docs.append(doc_data)
                else:
                    failed_files.append(file_info['file_path'])
                    
            except Exception as e:
                logger.error(f"❌ Failed to process {file_info['file_path']}: {e}")
                failed_files.append(file_info['file_path'])
        
        logger.info(f"✅ Processing complete! {len(processed_docs)} successful, {len(failed_files)} failed")
        
        if failed_files:
            logger.warning(f"Failed files: {failed_files[:5]}...")  # Show first 5
        
        self.processed_documents = processed_docs
        return processed_docs
    
    def process_single_file(self, file_path: str) -> Optional[Dict]:
        """
        Process a single file based on its extension
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            logger.warning(f"⚠️ Unsupported file type: {extension}")
            return None
        
        try:
            # Get the appropriate processor
            processor = self.supported_formats[extension]
            
            # Extract text content
            text_content = processor(file_path)
            
            if not text_content or len(text_content.strip()) < 10:
                logger.warning(f"⚠️ No meaningful content extracted from {file_path.name}")
                return None
            
            # Extract metadata and legal information
            metadata = self._extract_legal_metadata(text_content, file_path)
            
            # Generate embeddings if model available
            embeddings = None
            if self.sentence_model:
                try:
                    embeddings = self.sentence_model.encode(text_content[:1000])  # First 1000 chars
                except Exception as e:
                    logger.warning(f"⚠️ Could not generate embeddings for {file_path.name}: {e}")
            
            # Create document object
            doc_data = {
                'text': text_content,
                'metadata': metadata,
                'embeddings': embeddings.tolist() if embeddings is not None else None,
                'processed_time': datetime.now().isoformat(),
                'text_length': len(text_content),
                'chunk_count': len(self._chunk_text(text_content))
            }
            
            return doc_data
            
        except Exception as e:
            logger.error(f"❌ Error processing {file_path}: {e}")
            return None
    
    def _process_text(self, file_path: Path) -> str:
        """Process plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            logger.warning(f"⚠️ Could not decode {file_path}")
            return ""
    
    def _process_pdf(self, file_path: Path) -> str:
        """Process PDF files with multiple extraction methods"""
        text_content = ""
        
        # Method 1: pdfplumber (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages[:10]:  # Limit to first 10 pages
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
        except Exception as e:
            logger.warning(f"⚠️ pdfplumber failed for {file_path}: {e}")
        
        # Method 2: PyPDF2 as fallback
        if not text_content:
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num in range(min(10, len(pdf_reader.pages))):
                        page = pdf_reader.pages[page_num]
                        text_content += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"⚠️ PyPDF2 failed for {file_path}: {e}")
        
        # Method 3: textract as last resort
        if not text_content:
            try:
                text_content = textract.process(file_path).decode('utf-8')
            except Exception as e:
                logger.warning(f"⚠️ textract failed for {file_path}: {e}")
        
        return text_content
    
    def _process_json(self, file_path: Path) -> str:
        """Process JSON files and extract text content"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Extract text from various JSON structures
            text_parts = []
            
            def extract_text_recursive(obj, depth=0):
                if depth > 5:  # Prevent infinite recursion
                    return
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, str) and len(value) > 10:
                            text_parts.append(f"{key}: {value}")
                        elif isinstance(value, (dict, list)):
                            extract_text_recursive(value, depth + 1)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_text_recursive(item, depth + 1)
                elif isinstance(obj, str) and len(obj) > 10:
                    text_parts.append(obj)
            
            extract_text_recursive(data)
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"❌ Error processing JSON {file_path}: {e}")
            return ""
    
    def _process_docx(self, file_path: Path) -> str:
        """Process Word DOCX files"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"❌ Error processing DOCX {file_path}: {e}")
            return ""
    
    def _process_doc_fallback(self, file_path: Path) -> str:
        """Process legacy DOC files using textract"""
        try:
            return textract.process(file_path).decode('utf-8')
        except Exception as e:
            logger.error(f"❌ Error processing DOC {file_path}: {e}")
            return ""
    
    def _process_pages(self, file_path: Path) -> str:
        """Process Apple Pages files"""
        try:
            # Pages files are actually zip packages
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Look for the main content file
                for file_name in zip_file.namelist():
                    if 'index.xml' in file_name or 'content.xml' in file_name:
                        with zip_file.open(file_name) as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            
                            # Extract text from XML elements
                            text_parts = []
                            for elem in root.iter():
                                if elem.text and elem.text.strip():
                                    text_parts.append(elem.text.strip())
                            
                            return "\n".join(text_parts)
            
            # Fallback to textract
            return textract.process(file_path).decode('utf-8')
            
        except Exception as e:
            logger.error(f"❌ Error processing Pages {file_path}: {e}")
            return ""
    
    def _process_rtfd(self, file_path: Path) -> str:
        """Process RTFD (Rich Text Format Directory) files"""
        try:
            # RTFD is a package, look for TXT.rtf inside
            rtfd_path = Path(file_path)
            txt_rtf_path = rtfd_path / "TXT.rtf"
            
            if txt_rtf_path.exists():
                return textract.process(txt_rtf_path).decode('utf-8')
            else:
                # Try to find any .rtf file in the package
                for rtf_file in rtfd_path.glob("*.rtf"):
                    return textract.process(rtf_file).decode('utf-8')
            
            return ""
            
        except Exception as e:
            logger.error(f"❌ Error processing RTFD {file_path}: {e}")
            return ""
    
    def _process_rtf(self, file_path: Path) -> str:
        """Process RTF files"""
        try:
            return textract.process(file_path).decode('utf-8')
        except Exception as e:
            logger.error(f"❌ Error processing RTF {file_path}: {e}")
            return ""
    
    def _extract_category_from_path(self, relative_path: Path) -> str:
        """Extract document category from file path"""
        path_parts = relative_path.parts
        
        # Category mapping based on directory names
        category_mappings = {
            'DIV-COMMS': 'Communications',
            'DIV-MARIA TESTA': 'Attorney Communications',
            'DIV-PTSD-ADHD': 'Medical Records',
            'DIV-FMLA': 'FMLA Documentation',
            'DIV-TIMELINE OF EVENTS': 'Timeline',
            'DIV-PREC-CASES': 'Precedent Cases',
            'DIV-CONDO': 'Property Records',
            'DIV-BRIDGE LOAN': 'Financial Records',
            'DIV-MB-DISCLOSURE': 'Disclosure Documents',
            'SB-CASE-FILE': 'Case Files',
            'DIV-HELEN-OC-COMMS': 'Opposing Counsel Communications'
        }
        
        for part in path_parts:
            for key, category in category_mappings.items():
                if key in part.upper():
                    return category
        
        # Default categorization
        if 'EMAIL' in str(relative_path).upper():
            return 'Email Communications'
        elif 'CASE' in str(relative_path).upper():
            return 'Case Documents'
        elif 'FINANCIAL' in str(relative_path).upper():
            return 'Financial Records'
        else:
            return 'General Documents'
    
    def _extract_legal_metadata(self, text: str, file_path: Path) -> Dict:
        """Extract legal-specific metadata from document text"""
        metadata = {
            'legal_claims': [],
            'case_citations': [],
            'statute_citations': [],
            'court_rules': [],
            'dates': [],
            'dollar_amounts': [],
            'legal_entities': [],
            'potential_misinformation': []
        }
        
        # Extract patterns
        for pattern_name, pattern in self.legal_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                metadata[pattern_name] = list(set(matches))  # Remove duplicates
        
        # Extract legal claims for fact-checking
        for pattern in self.legal_claim_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            metadata['legal_claims'].extend(matches)
        
        # Extract names (potential legal entities)
        name_pattern = r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b'
        names = re.findall(name_pattern, text)
        metadata['legal_entities'] = list(set(names))[:10]  # Top 10 unique names
        
        # Check for potential misinformation keywords
        misinformation_keywords = [
            'financial abandonment doctrine',
            'abandonment doctrine',
            'medical abandonment doctrine'
        ]
        
        for keyword in misinformation_keywords:
            if keyword.lower() in text.lower():
                metadata['potential_misinformation'].append(keyword)
        
        return metadata
    
    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks for better semantic search"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence ending within the last 100 characters
                sentence_end = text.rfind('.', start + chunk_size - 100, end)
                if sentence_end != -1:
                    end = sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file for change detection"""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            logger.warning(f"⚠️ Could not calculate hash for {file_path}: {e}")
            return ""
    
    def save_processed_documents(self, output_file: str = "processed_documents.json") -> None:
        """Save processed documents to JSON file"""
        try:
            # Convert numpy arrays to lists for JSON serialization
            documents_for_json = []
            for doc in self.processed_documents:
                doc_copy = doc.copy()
                if doc_copy.get('embeddings'):
                    doc_copy['embeddings'] = doc_copy['embeddings']  # Already converted to list
                documents_for_json.append(doc_copy)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(documents_for_json, f, indent=2, ensure_ascii=False)
            
            logger.info(f"✅ Saved {len(documents_for_json)} processed documents to {output_file}")
            
        except Exception as e:
            logger.error(f"❌ Error saving processed documents: {e}")
    
    def get_processing_summary(self) -> Dict:
        """Get summary of processing results"""
        if not self.processed_documents:
            return {"status": "No documents processed"}
        
        total_docs = len(self.processed_documents)
        total_text_length = sum(doc.get('text_length', 0) for doc in self.processed_documents)
        
        # Count by category
        categories = {}
        extensions = {}
        
        for doc in self.processed_documents:
            category = doc.get('category', 'Unknown')
            extension = doc.get('extension', 'Unknown')
            
            categories[category] = categories.get(category, 0) + 1
            extensions[extension] = extensions.get(extension, 0) + 1
        
        return {
            "total_documents": total_docs,
            "total_text_length": total_text_length,
            "average_text_length": total_text_length / total_docs if total_docs > 0 else 0,
            "categories": categories,
            "file_extensions": extensions,
            "processing_time": datetime.now().isoformat(),
            "has_embeddings": any(doc.get('embeddings') for doc in self.processed_documents)
        }

# Test the processor if run directly
if __name__ == "__main__":
    processor = ComprehensiveDocumentProcessor()
    
    print("🧪 Testing Comprehensive Document Processor...")
    
    # Test 1: Scan directories
    print("\n1. Scanning directories...")
    files = processor.scan_all_directories()
    print(f"Found {len(files)} files")
    
    # Test 2: Process a few documents
    print("\n2. Processing first 5 documents...")
    processed = processor.process_all_documents(max_files=5)
    
    # Test 3: Get summary
    print("\n3. Processing Summary:")
    summary = processor.get_processing_summary()
    for key, value in summary.items():
        print(f"{key}: {value}")
    
    # Test 4: Save results
    print("\n4. Saving processed documents...")
    processor.save_processed_documents("test_processed_documents.json")
    
    print("\n✅ Document processor testing completed!")